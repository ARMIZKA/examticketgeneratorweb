import random
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from jinja2 import Template
import subprocess
import os
from jinja2 import Environment, BaseLoader

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "2")  # тонкие линии
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def read_questions_from_docx(path):
    doc = DocxDocument(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def simple_pairing(questions, num_tickets):
    if len(questions) < num_tickets * 2:
        raise ValueError("Недостаточно вопросов для формирования билетов.")
    random.shuffle(questions)
    needed = questions[:num_tickets*2]
    return [needed[i*2:(i+1)*2] for i in range(num_tickets)]

def create_formatted_exam_docx(tickets, output_path,
                               discipline="Операционные системы",
                               specialty="02.03.02 ФИиИТ",
                               group="БФИ2102",
                               teacher="И.О. Фамилия",
                               deputy="П.П. Заместитель"):
    doc = DocxDocument()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    for i, ticket in enumerate(tickets, 1):
        header_table = doc.add_table(rows=1, cols=3)
        set_table_borders(header_table)
        widths = [Cm(6), Cm(6), Cm(6)]
        for idx, cell in enumerate(header_table.rows[0].cells):
            cell.width = widths[idx]
        header_table.cell(0, 0).text = (
            "Рассмотрено ПЦК\n"
            "Протокол № ___ от ______\n"
            "__________________ Ф.И.О.\n"
            "председатель ПЦК"
        )
        header_table.cell(0, 1).text = (
            f"Экзаменационный билет № {i}\n"
            f"Дисциплина: {discipline}\n"
            f"Специальность: {specialty}\n"
            f"Группа: {group}"
        )
        header_table.cell(0, 2).text = (
            "УТВЕРЖДАЮ\n"
            "Заместитель директора по УВР\n"
            "____________ " + deputy + "\n"
            "«____» ___________ 20__ год"
        )
        for cell in header_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph("1. " + ticket[0])
        doc.add_paragraph("2. " + ticket[1])
        doc.add_paragraph("3.* " + "______________________________________________________________")
        doc.add_paragraph()
        p_teacher = doc.add_paragraph("Преподаватель ___________________ " + teacher)
        p_teacher.paragraph_format.left_indent = Cm(0)
        p_teacher.paragraph_format.first_line_indent = Cm(0)
        doc.add_page_break()

    doc.save(output_path)

latex_template = r"""
\documentclass[8pt,a4paper]{article}
\usepackage{fontspec} 
\setmainfont{Times New Roman} 
\usepackage[russian]{babel}
\usepackage{geometry}
\geometry{left=2cm,right=2cm,top=2cm,bottom=2cm}
\usepackage{graphicx}
\usepackage{setspace}
\usepackage{enumitem}
\usepackage{ragged2e}
\usepackage{titlesec}
\setlist[enumerate]{label=\arabic*., leftmargin=*}
\pagestyle{empty}

\begin{document}

{% for ticket in tickets %}
{% raw %}
\noindent
\fbox{%
    \parbox[t][0.48\textheight][t]{\textwidth}{%
        \fontsize{6pt}{6.5pt}\selectfont
        \setstretch{1}
        \begin{minipage}[t]{0.45\textwidth}
            \vspace{-0.2pt}
            \centering
            \includegraphics[height=3\baselineskip]{D:/logo.pdf}\\[0.1cm]
            {\scriptsize
            МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ, СВЯЗИ И МАССОВЫХ\\
            КОММУНИКАЦИЙ РОССИЙСКОЙ ФЕДЕРАЦИИ\\
            Ордена Трудового Красного Знамени федеральное государственное\\
            бюджетное образовательное учреждение высшего образования\\
            «Московский технический университет связи и информатики» (МТУСИ)
            }
        \end{minipage}%
        \hfill
        \begin{minipage}[t]{0.48\textwidth}
            \vspace{-0.2pt}
            {\small
            Направление подготовки: {% endraw %}{{ direction }}{% raw %}\\
            Профили подготовки: {% endraw %}{{ profile }}{% raw %}\\
            Кафедра: {% endraw %}{{ department }}{% raw %}\\
            Дисциплина: {% endraw %}{{ discipline }}{% raw %}
            }
        \end{minipage}

        \vspace{0.4cm}

        \begin{center}
            {\normalsize \textbf{ЭКЗАМЕНАЦИОННЫЙ БИЛЕТ № {% endraw %}{{ loop.index }}{% raw %}}}
        \end{center}

        \vspace{0.3cm}
        \fontsize{8pt}{9pt}\selectfont

        \begin{enumerate}
            \item {% endraw %}{{ ticket[0] }}{% raw %}
            \item {% endraw %}{{ ticket[1] }}{% raw %}
            \item Практическое задание (выдается преподавателем).
        \end{enumerate}

        \vspace*{\fill}

Составитель \hspace{5.1cm} \underline{\hspace{4cm}} \hfill к.т.н., доцент Ковтун И.И.\\[0.3cm]
Зав. кафедрой \hspace{5cm} \underline{\hspace{4cm}} \hfill д.т.н., профессор Полищук Ю.В.\\[0.3cm]
\begin{center}
    {\footnotesize Утверждено на заседании кафедры № 2 от «03» октября 2024 г.}
\end{center}

    }
}
\newpage
{% endraw %}
{% endfor %}

\end{document}
"""

def generate_latex(tickets, output_file, direction, profile, department, discipline):
    env = Environment(loader=BaseLoader(), autoescape=False)
    template = env.from_string(latex_template)
    rendered = template.render(
        tickets=tickets,
        direction=direction,
        profile=profile,
        department=department,
        discipline=discipline
    )
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(rendered)
def compile_pdf(tex_file, output_dir):
    try:
        pdflatex_path = r"C:\Users\artem\AppData\Local\Programs\MiKTeX\miktex\bin\x64\xelatex.exe"
        subprocess.run([pdflatex_path, "-output-directory", output_dir, tex_file], check=True)
        return True
    except subprocess.CalledProcessError:
        return False

def generate_tickets_from_questions(docx_path):

    questions = read_questions_from_docx(docx_path)
    tickets = []
    i = 0
    while i < len(questions):
        q1 = questions[i]
        q2 = questions[i + 1] if (i + 1) < len(questions) else ""
        tickets.append([q1, q2])
        i += 2
    return tickets

